"""Extract and chunk local documents before embedding them."""

import os
import re
from pathlib import Path

from embedding import EMBEDDER


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


def clean_text(value: object) -> str:
    """Replace invalid lone surrogate code points from extracted documents."""
    return "".join(
        "\ufffd" if 0xD800 <= ord(character) <= 0xDFFF else character
        for character in str(value)
    )


def normalize(text: str) -> str:
    paragraphs = []
    for paragraph in re.split(
        r"\n\s*\n",
        clean_text(text).replace("\x00", " "),
    ):
        cleaned = re.sub(r"\s+", " ", paragraph).strip()
        if cleaned:
            paragraphs.append(cleaned)
    return "\n\n".join(paragraphs)


def extract(path: Path) -> str:
    extension = path.suffix.lower()
    if extension == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as error:
            raise RuntimeError("pypdf is required for PDF documents") from error
        text = "\n\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
    elif extension == ".docx":
        try:
            from docx import Document
        except ImportError as error:
            raise RuntimeError("python-docx is required for DOCX documents") from error
        document = Document(path)
        blocks = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.append(" | ".join(cell.text for cell in row.cells))
        text = "\n\n".join(blocks)
    elif extension in {".txt", ".md"}:
        try:
            text = path.read_text(encoding="utf-8-sig")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")
    else:
        raise ValueError(
            "Unsupported document type. Allowed: "
            + ", ".join(sorted(SUPPORTED_EXTENSIONS))
        )

    text = normalize(text)
    if not text:
        raise ValueError("The document contains no extractable text")
    return text


def chunk(text: str) -> list[str]:
    size = max(40, int(os.getenv("RAGDOLL_CHUNK_WORDS", "180")))
    overlap = max(0, int(os.getenv("RAGDOLL_CHUNK_OVERLAP_WORDS", "30")))
    overlap = min(overlap, size - 1)
    words = normalize(text).split()
    return [
        " ".join(words[start : start + size])
        for start in range(0, len(words), size - overlap)
        if words[start : start + size]
    ]


def prepare(path_value: str) -> dict[str, object]:
    path = Path(path_value).resolve()
    chunks = chunk(extract(path))
    vectors = EMBEDDER.encode(chunks)
    return {
        "chunks": chunks,
        "embeddings": vectors,
        "embedding_dimension": len(vectors[0]) if vectors else 0,
        "embedding_model": EMBEDDER.model_name,
    }