import os
import sys

# llama setup stuff
os.environ.setdefault("PYTHONDONTWRITEBYTECODE", "1")
sys.dont_write_bytecode = True

import re
import threading
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, List, Sequence

try:
    import numpy as np
except ImportError:  # an error is raised when embeddings are requested.
    np = None  # type: ignore[assignment]

try:
    import torch
except ImportError:
    torch = None  # type: ignore[assignment]

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None  # type: ignore[assignment,misc]


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_MODEL_DIRECTORY = PROJECT_ROOT / "models" / "embedding" / "bge-base-en-v1.5"
EMBEDDING_MODEL_DIRECTORY = Path(
    os.getenv("RAGDOLL_EMBEDDING_MODEL_DIR", str(DEFAULT_MODEL_DIRECTORY))
).expanduser().resolve()
EMBEDDING_MODEL_NAME = (
    os.getenv("RAGDOLL_EMBEDDING_MODEL_NAME", "BAAI/bge-base-en-v1.5").strip()
    or "BAAI/bge-base-en-v1.5"
)

DEFAULT_QUERY_PREFIX = "Represent this sentence for searching relevant passages: "
QUERY_PREFIX = os.getenv("RAGDOLL_EMBEDDING_QUERY_PREFIX", DEFAULT_QUERY_PREFIX)
CHUNK_WORDS = max(40, int(os.getenv("RAGDOLL_CHUNK_WORDS", "180")))
CHUNK_OVERLAP_WORDS = max(0, int(os.getenv("RAGDOLL_CHUNK_OVERLAP_WORDS", "30")))
EMBEDDING_BATCH_SIZE = max(1, int(os.getenv("RAGDOLL_EMBEDDING_BATCH_SIZE", "16")))

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


class EmbeddingConfigurationError(RuntimeError):
    """Raised when the local embedding runtime is not ready."""


class DocumentExtractionError(RuntimeError):
    """Raised when text cannot be extracted from an uploaded document."""


@dataclass(frozen=True)
class PreparedDocument:
    text: str
    chunks: List[str]
    embeddings: List[List[float]]
    embedding_dimension: int
    embedding_model_name: str
    model_directory: str


class LocalSentenceEmbedder:
    """Lazy, thread-safe wrapper around a local SentenceTransformer model."""

    def __init__(self, model_directory: Path = EMBEDDING_MODEL_DIRECTORY) -> None:
        self.model_directory = model_directory
        self._model: Any = None
        self._load_lock = threading.Lock()
        self._encode_lock = threading.Lock()

    def _resolve_device(self) -> str:
        configured = os.getenv("RAGDOLL_EMBEDDING_DEVICE", "").strip()
        if configured:
            return configured
        if torch is not None and bool(torch.cuda.is_available()):
            return "cuda"
        return "cpu"

    def _model_kwargs(self) -> dict[str, Any]:
        requested_dtype = os.getenv("RAGDOLL_EMBEDDING_DTYPE", "").strip().lower()
        if not requested_dtype or torch is None:
            return {}

        supported = {
            "float16": torch.float16,
            "fp16": torch.float16,
            "float32": torch.float32,
            "fp32": torch.float32,
            "bfloat16": torch.bfloat16,
            "bf16": torch.bfloat16,
        }
        if requested_dtype not in supported:
            raise EmbeddingConfigurationError(
                "RAGDOLL_EMBEDDING_DTYPE must be float16, float32, or bfloat16."
            )
        return {"torch_dtype": supported[requested_dtype]}

    def _validate_model_directory(self) -> None:
        if not self.model_directory.is_dir():
            raise EmbeddingConfigurationError(
                "The local embedding model directory was not found: "
                f"{self.model_directory}. Download the complete "
                "BAAI/bge-base-en-v1.5 repository into that folder or set "
                "RAGDOLL_EMBEDDING_MODEL_DIR."
            )
        if not (self.model_directory / "config.json").is_file():
            raise EmbeddingConfigurationError(
                f"{self.model_directory} is missing config.json. Download the "
                "complete Hugging Face repository, not only model.safetensors."
            )

    def get_model(self) -> Any:
        if self._model is not None:
            return self._model

        with self._load_lock:
            if self._model is not None:
                return self._model
            if SentenceTransformer is None:
                raise EmbeddingConfigurationError(
                    "sentence-transformers is not installed. Run "
                    "python -m pip install -r requirements.txt."
                )
            if np is None:
                raise EmbeddingConfigurationError(
                    "numpy is not installed. Run python -m pip install -r requirements.txt."
                )

            self._validate_model_directory()
            self._model = SentenceTransformer(
                str(self.model_directory),
                device=self._resolve_device(),
                local_files_only=True,
                model_kwargs=self._model_kwargs(),
            )
            return self._model

    def encode(self, texts: Sequence[str], *, is_query: bool = False) -> List[List[float]]:
        cleaned = [normalize_text(text) for text in texts]
        if any(not text for text in cleaned):
            raise ValueError("Embedding input cannot be empty.")
        if is_query and QUERY_PREFIX:
            cleaned = [QUERY_PREFIX + text for text in cleaned]

        model = self.get_model()
        with self._encode_lock:
            vectors = model.encode(
                cleaned,
                batch_size=EMBEDDING_BATCH_SIZE,
                convert_to_numpy=True,
                normalize_embeddings=True,
                show_progress_bar=False,
            )
        return vectors.astype("float32").tolist()

    def status(self) -> dict[str, Any]:
        directory_exists = self.model_directory.is_dir()
        model_loaded = self._model is not None
        return {
            "model_name": EMBEDDING_MODEL_NAME,
            "model_directory": str(self.model_directory),
            "directory_exists": directory_exists,
            "config_present": (self.model_directory / "config.json").is_file(),
            "model_loaded": model_loaded,
            "device": self._resolve_device(),
            "offline_only": True,
            "supported_extensions": sorted(SUPPORTED_EXTENSIONS),
        }


EMBEDDER = LocalSentenceEmbedder()


def normalize_text(text: str) -> str:
    """Normalize whitespace while retaining paragraph boundaries."""

    text = str(text or "").replace("\x00", " ").replace("\r\n", "\n")
    paragraphs = []
    for paragraph in re.split(r"\n\s*\n", text):
        cleaned = re.sub(r"[\t ]+", " ", paragraph)
        cleaned = re.sub(r"\s*\n\s*", " ", cleaned).strip()
        if cleaned:
            paragraphs.append(cleaned)
    return "\n\n".join(paragraphs)


def chunk_text(
    text: str,
    *,
    chunk_words: int = CHUNK_WORDS,
    overlap_words: int = CHUNK_OVERLAP_WORDS,
) -> List[str]:
    """Split text into overlapping, retrieval-sized word chunks."""

    normalized = normalize_text(text)
    words = normalized.split()
    if not words:
        return []

    chunk_words = max(20, chunk_words)
    overlap_words = min(max(0, overlap_words), chunk_words - 1)
    step = chunk_words - overlap_words

    chunks: List[str] = []
    for start in range(0, len(words), step):
        chunk = " ".join(words[start : start + chunk_words]).strip()
        if chunk:
            chunks.append(chunk)
        if start + chunk_words >= len(words):
            break
    return chunks


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise DocumentExtractionError(
            "pypdf is required for PDF uploads. Install requirements.txt."
        ) from error

    try:
        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as error:
        raise DocumentExtractionError(f"Unable to read PDF: {error}") from error
    return "\n\n".join(page for page in pages if page)


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as error:
        raise DocumentExtractionError(
            "python-docx is required for DOCX uploads. Install requirements.txt."
        ) from error

    try:
        document = Document(str(path))
        blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
    except Exception as error:
        raise DocumentExtractionError(f"Unable to read DOCX: {error}") from error
    return "\n\n".join(blocks)


def extract_document_text(path: Path) -> str:
    """Extract normalized text from a supported local document."""

    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DocumentExtractionError(
            f"Unsupported document type {extension or '(none)'}. Allowed: {allowed}."
        )

    if extension == ".pdf":
        text = _extract_pdf(path)
    elif extension == ".docx":
        text = _extract_docx(path)
    else:
        try:
            text = path.read_text(encoding="utf-8-sig")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")
        except OSError as error:
            raise DocumentExtractionError(f"Unable to read document: {error}") from error

    normalized = normalize_text(text)
    if not normalized:
        raise DocumentExtractionError(
            "No extractable text was found. Scanned/image-only PDFs require OCR, "
            "which is not enabled in this project."
        )
    return normalized


def embed_documents(texts: Sequence[str]) -> List[List[float]]:
    return EMBEDDER.encode(texts, is_query=False)


def embed_query(text: str) -> List[float]:
    return EMBEDDER.encode([text], is_query=True)[0]


def prepare_document(path: Path) -> PreparedDocument:
    text = extract_document_text(path)
    chunks = chunk_text(text)
    if not chunks:
        raise DocumentExtractionError("The document did not produce any chunks.")
    embeddings = embed_documents(chunks)
    dimension = len(embeddings[0]) if embeddings else 0
    return PreparedDocument(
        text=text,
        chunks=chunks,
        embeddings=embeddings,
        embedding_dimension=dimension,
        embedding_model_name=EMBEDDING_MODEL_NAME,
        model_directory=str(EMBEDDER.model_directory),
    )